from dotenv import load_dotenv
import os
from langchain_openai import OpenAI
from langchain_core.prompts import PromptTemplate
from langchain.schema import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
import pdfplumber
from docx import Document as DocxDocument
import heapq

class LLMQuery:
    def __init__(self):
        load_dotenv()
        openai_api_key = os.getenv("OPENAI_API_KEY")
        self.llm = OpenAI(api_key=openai_api_key)
        self.embeddings = OpenAIEmbeddings(api_key=openai_api_key)

    def _load_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        """
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def _load_word(self, file_path: str) -> str:
        """
        Extract text from a Word file.
        """
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def _load_documents(self, file_paths: list[str]) -> list[Document]:
        """
        Load documents from a list of file paths (PDF and Word).
        """
        documents = []
        for file_path in file_paths:
            if file_path.lower().endswith('.pdf'):
                content = self._load_pdf(file_path)
            elif file_path.lower().endswith('.docx'):
                content = self._load_word(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path}")

            documents.append(Document(page_content=content))
        return documents

    def _split_and_embed_documents(self, documents: list[Document]) -> FAISS:
        """
        Split documents into smaller chunks and embed them into a vector store.
        """
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks = [text_splitter.split_text(doc.page_content) for doc in documents]
        flat_chunks = [chunk for sublist in chunks for chunk in sublist]
        doc_vectors = FAISS.from_texts(flat_chunks, self.embeddings)
        return doc_vectors

    def _split_context(self, context: str, max_chunk_size: int = 1500) -> list[str]:
        """
        Split the context into smaller chunks to avoid exceeding the token limit.

        :param context: The full context text to split.
        :param max_chunk_size: Maximum allowed size for each chunk in words.
        :return: A list of context chunks.
        """
        words = context.split()
        chunks = [words[i:i + max_chunk_size] for i in range(0, len(words), max_chunk_size)]
        return [" ".join(chunk) for chunk in chunks]

    def _estimate_token_count(self, text: str) -> int:
        """
        Estimate the token count of a given text.
        """
        return len(text) // 4

    def _rank_sentences(self, responses: list[str], query: str) -> str:
        """
        Rank sentences based on relevance to the query by passing the sentences to the LLM 
        with the query and a prompt to assess their fit.

        :param responses: A list of response texts.
        :param query: The user query for context.
        :return: A single string with the top sentences ranked by relevance.
        """
        all_sentences = []
        for response in responses:
            sentences = response.split('. ')
            all_sentences.extend(sentences)

        ranking_prompt_template = (
            "Query:\n{query}\n\n"
            "Sentence:\n{sentence}\n\n"
            "Rate the relevance of the sentence to the query on a scale of 1 to 10, "
            "where 10 means highly relevant and 1 means not relevant at all. Just provide the rating."
        )
        ranking_prompt = PromptTemplate(
            template=ranking_prompt_template,
            input_variables=["query", "sentence"]
        )

        scored_sentences = []
        for sentence in all_sentences:
            if not sentence.strip():  # incase there are any empty sentences
                continue

           
            chain = ranking_prompt | self.llm
            rating_response = chain.invoke({"query": query, "sentence": sentence})

            # parsing the response as an integer score
            try:
                rating = int(rating_response.strip())
            except ValueError:
                rating = 0  # if parsing fails

            scored_sentences.append((rating, sentence))

        scored_sentences.sort(reverse=True, key=lambda x: x[0])

        # Top 3 sentences for now
        top_sentences = [sentence for _, sentence in scored_sentences[:3]]

        return '.\n'.join(top_sentences) + '.\n'




    def generate_query(self, file_paths: list[str], few_shot_prompts: list[str], query: str) -> str:
        """
        Generate a query using the provided documents (PDF and Word files), few-shot prompts, and user query.

        :param file_paths: List of file paths to the documents. If empty, only few-shot prompts and query will be used.
        :param few_shot_prompts: List of few-shot examples to guide the model.
        :param query: The user query for which we want to generate a response.
        :return: The response generated by the LLM.
        """
        context = ""

        
        if file_paths:
            documents = self._load_documents(file_paths)
            if documents:
                doc_vectors = self._split_and_embed_documents(documents)
                retrieved_chunks: list[Document] = doc_vectors.similarity_search(query, k=1)
                context = "\n".join([chunk.page_content for chunk in retrieved_chunks])

                
                context_chunks = self._split_context(context, max_chunk_size=1500)
            else:
                print("Failed to load content from the provided documents.")
        else:
            context_chunks = [""]


        prompt_template = "\n".join(few_shot_prompts) + "\n\nContext:\n{context}\n\nQuery:\n{query}\n\nAnswer:"
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "query"]
        )

        responses = []

        
        for chunk in context_chunks:
            estimated_tokens = self._estimate_token_count(chunk + query + "".join(few_shot_prompts))
            
            if estimated_tokens > 4000:
                sub_chunks = self._split_context(chunk, max_chunk_size=1000)
                partial_response = ""
                
                for sub_chunk in sub_chunks:
                    estimated_sub_tokens = self._estimate_token_count(sub_chunk + query + "".join(few_shot_prompts))
                    
                    if estimated_sub_tokens > 4000:
                        print(f"Skipping sub-chunk due to exceeding token limit: {estimated_sub_tokens} tokens.")
                        continue
                    
                    chain = prompt | self.llm
                    response = chain.invoke({"context": sub_chunk, "query": query})
                    responses.append(response)
                    partial_response += response + " "
                
                
                responses.append(partial_response.strip())
            else:
                # If the chunk is within the token limit, process it as usual.
                chain = prompt | self.llm
                response = chain.invoke({"context": chunk, "query": query})
                responses.append(response)

        # Getting the top 3 sentences from each prompt
        combined_response = self._rank_sentences(responses, query)
        return combined_response


if __name__ == "__main__":
    llm_query_rag = LLMQuery()
    print(llm_query_rag.generate_query(
        ["tst/test_doc.docx"], 
        [], 
        "Who was the founder of the company 'test', and when was it founded?"
    ))
